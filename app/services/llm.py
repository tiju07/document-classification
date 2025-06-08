import os
import mimetypes
from typing import List
from openai import AzureOpenAI
from PIL import Image
from io import BytesIO
import base64
from docx import Document
import fitz  # PyMuPDF

llm = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_KEY"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
)

def encode_image_to_base64(image: Image.Image) -> str:
    buffered = BytesIO()
    image.save(buffered, format="PNG")
    img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
    return f"data:image/png;base64,{img_str}"

def extract_entities(text: str) -> str:
    # return '{\"key\": \"LLM Response\"}'
    response = llm.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", 
             "content": """You are an expert in extracting structured entities from email content.
                            Given the body of an email, extract the following entities:
                            Dates: All relevant dates mentioned.
                            Parties: Names of individuals, companies, or organizations.
                            Amounts: All monetary values with currency symbols.
                            Other Entities: Any additional relevant information (like Invoice No, VAT No, IBAN, BIC, etc.)
                            Instructions:
                            Return the output strictly as a valid JSON object (no Markdown code blocks, no explanations, no comments).
                            Do not include any text before or after the JSON.
                            Use double quotes (") for all JSON keys and string values.
                            Use this structure:
                            {
                                "Dates": [],
                                "Parties": [],
                                "Amounts": [],
                                "Other Entities": {
                                    // keys and values dynamically extracted from the email
                                }
                            }"""},
            {"role": "user", "content": text}
        ],
        max_tokens=500
    )
    return response.choices[0].message.content.strip()

def classify_document(text: str, entities: str = None) -> str:
    # return '{\"category\": \"contract\", \"confidence_score\":80}'
    messages = [
        {"role": "system", 
        "content": """You are an expert in extracting structured entities from email content.
            Given the body of an email, extract the following entities:
            category: The category of the given content (e.g., contract, invoice, letter).
            confidence_score: The confidence score of the classification (0-100).
            Instructions:
            Return the output strictly as a valid JSON object (no Markdown code blocks, no explanations, no comments).
            Do not include any text before or after the JSON.
            Use double quotes (") for all JSON keys and string values.
            Use this structure:
            {
                "category": "",
                "confidence_score": ""
            }
        """},
        {"role": "user", "content": "Text: " + text + "\nEntities: " + entities}
    ]
    # if entities:
    #     messages.append({"role": "user", "content": entities})
    response = llm.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        max_tokens=500
    )
    return response.choices[0].message.content.strip()

def summarize_email_body(text: str) -> str:
    return text

def summarize_image_with_vision(image: Image.Image) -> str:
    # return "LLM Response"
    encoded_image = encode_image_to_base64(image)
    response = llm.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a document analysis assistant. Explain what's in this image."},
            {"role": "user", "content": [{"type": "image_url", "image_url": {"url": encoded_image}}]}
        ],
        max_tokens=500
    )
    return response.choices[0].message.content.strip()

def extract_images_from_pdf(pdf_path: str):
    images = []
    doc = fitz.open(pdf_path)
    for page in doc:
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            image = Image.open(BytesIO(img_bytes))
            images.append(image)
    return images

def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    doc = fitz.open(pdf_path)
    for page in doc:
        text += page.get_text()
    return text.strip()

def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs]).strip()

def extract_images_from_docx(docx_path: str):
    images = []
    doc = Document(docx_path)
    for rel in doc.part._rels:
        rel = doc.part._rels[rel]
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image = Image.open(BytesIO(image_data))
            images.append(image)
    return images

def is_image_file(file_path: str) -> bool:
    mime = mimetypes.guess_type(file_path)[0]
    return mime and mime.startswith("image")

def process_file(file_path: str) -> str:
    if is_image_file(file_path):
        print("Detected image file.")
        img = Image.open(file_path)
        return summarize_image_with_vision(img)

    images = []
    extracted_text = ""

    if file_path.endswith(".pdf"):
        print("Detected PDF file.")
        extracted_text = extract_text_from_pdf(file_path)
        images = extract_images_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        print("Detected DOCX file.")
        extracted_text = extract_text_from_docx(file_path)
        images = extract_images_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type")

    summaries = []

    if extracted_text:
        print("Extracted text...")
        summaries.append(extracted_text)

    for idx, image in enumerate(images):
        print(f"Summarizing image {idx+1}...")
        summary = summarize_image_with_vision(image)
        summaries.append(summary)

    return "\n".join(summaries)

